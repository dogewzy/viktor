"""Web chat file upload, extraction, and downloadable artifact helpers."""

from __future__ import annotations

import csv
import io
import json
import mimetypes
import re
from datetime import datetime
from pathlib import Path
from typing import Any
from uuid import uuid4

import oss2

from settings import file_upload_config, oss_config


_SAFE_NAME_RE = re.compile(r"[^A-Za-z0-9._\-\u4e00-\u9fff]+")
_TEXT_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".html", ".htm", ".log", ".json", ".csv", ".tsv", ".yaml", ".yml",
    ".sql", ".xml", ".ini", ".cfg", ".properties", ".py", ".java", ".go", ".ts", ".js",
}
_EXPORT_KEYWORDS = ("导出", "excel", "word", "xlsx", "csv")


def is_export_request(text: str) -> bool:
    lower = (text or "").lower()
    return any(keyword in lower for keyword in _EXPORT_KEYWORDS)


def _ensure_oss_enabled() -> None:
    if not oss_config.enabled:
        raise RuntimeError("OSS 配置不完整，请配置 access_key / secret_key / end_point / bucket")


def _auth() -> oss2.Auth:
    _ensure_oss_enabled()
    return oss2.Auth(oss_config.access_key, oss_config.secret_key)


def _bucket(*, public: bool = False) -> oss2.Bucket:
    endpoint = oss_config.end_point_public if public else oss_config.end_point
    return oss2.Bucket(_auth(), endpoint, oss_config.bucket)


def _safe_filename(filename: str) -> str:
    name = Path(filename or "file").name.strip().strip(".")
    if not name:
        name = "file"
    safe = _SAFE_NAME_RE.sub("_", name)
    return safe[:180] or "file"


def _extension(filename: str) -> str:
    return Path(filename or "").suffix.lower()


def _validate_file(filename: str, size: int) -> None:
    if size <= 0:
        raise ValueError("文件为空，无法上传")
    if size > file_upload_config.max_file_size_bytes:
        raise ValueError(f"文件超过大小限制：{file_upload_config.max_file_size_mb}MB")
    ext = _extension(filename)
    if ext not in set(file_upload_config.allowed_extensions):
        supported = ", ".join(file_upload_config.allowed_extensions)
        raise ValueError(f"暂不支持的文件类型 {ext or '(无扩展名)'}；支持：{supported}")


def _object_key(project_id: str, topic_thread_id: str, filename: str, *, generated: bool = False) -> str:
    safe_project = _SAFE_NAME_RE.sub("_", project_id or "unknown")[:80] or "unknown"
    safe_topic = _SAFE_NAME_RE.sub("_", topic_thread_id or "generated")[:120] or "generated"
    bucket_dir = "generated" if generated else "uploads"
    date_part = datetime.now().strftime("%Y%m%d")
    return (
        f"{oss_config.upload_prefix}/{safe_project}/{safe_topic}/"
        f"{bucket_dir}/{date_part}/{uuid4().hex}/{_safe_filename(filename)}"
    )


def signed_download_url(object_key: str) -> str:
    return _bucket(public=True).sign_url(
        "GET",
        object_key,
        max(oss_config.signed_url_ttl_seconds, 60),
        slash_safe=True,
    )


def _oss_uri(object_key: str) -> str:
    return f"oss://{oss_config.bucket}/{object_key}"


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _truncate(text: str, max_chars: int) -> tuple[str, bool]:
    if len(text) <= max_chars:
        return text, False
    return text[:max_chars] + "\n\n... [内容过长，已截断] ...", True


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for i, page in enumerate(reader.pages[:50], start=1):
        pages.append(f"## Page {i}\n{page.extract_text() or ''}".strip())
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables[:20]:
        for row in table.rows[:100]:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    sections: list[str] = []
    for ws in wb.worksheets[:10]:
        rows: list[str] = [f"## Sheet: {ws.title}"]
        for row_index, row in enumerate(ws.iter_rows(values_only=True), start=1):
            if row_index > file_upload_config.max_excel_rows:
                rows.append(f"... [工作表超过 {file_upload_config.max_excel_rows} 行，已截断]")
                break
            values = ["" if cell is None else str(cell) for cell in row[: file_upload_config.max_excel_cols]]
            if any(v.strip() for v in values):
                rows.append("\t".join(values))
        sections.append("\n".join(rows))
    return "\n\n".join(sections)


def extract_file_text(data: bytes, filename: str) -> tuple[str, str, bool]:
    ext = _extension(filename)
    try:
        if ext in _TEXT_EXTENSIONS:
            text = _decode_text(data)
        elif ext == ".pdf":
            text = _extract_pdf(data)
        elif ext == ".docx":
            text = _extract_docx(data)
        elif ext == ".xlsx":
            text = _extract_xlsx(data)
        else:
            return "", "unsupported", False
        text, truncated = _truncate(text.strip(), file_upload_config.extracted_max_chars)
        return text, "ok" if text else "empty", truncated
    except Exception as e:  # noqa: BLE001
        return f"文件读取失败：{e}", "error", False


def upload_user_file(
    *,
    project_id: str,
    topic_thread_id: str,
    filename: str,
    content_type: str,
    data: bytes,
) -> dict[str, Any]:
    _validate_file(filename, len(data))
    safe_name = _safe_filename(filename)
    object_key = _object_key(project_id, topic_thread_id, safe_name)
    detected_type = content_type or mimetypes.guess_type(safe_name)[0] or "application/octet-stream"
    _bucket().put_object(object_key, data, headers={"Content-Type": detected_type})

    extracted_text, extract_status, truncated = extract_file_text(data, safe_name)
    preview, _ = _truncate(extracted_text, file_upload_config.preview_max_chars)
    return {
        "file_id": uuid4().hex,
        "filename": safe_name,
        "size": len(data),
        "content_type": detected_type,
        "oss_uri": _oss_uri(object_key),
        "object_key": object_key,
        "download_url": signed_download_url(object_key),
        "extract_status": extract_status,
        "extracted_text": extracted_text,
        "extracted_preview": preview,
        "truncated": truncated,
    }


def _csv_rows(content: str) -> list[list[str]]:
    sample = content.strip()
    if not sample:
        return [[""]]
    try:
        dialect = csv.Sniffer().sniff(sample[:2048])
    except csv.Error:
        dialect = csv.excel
    return [row for row in csv.reader(io.StringIO(sample), dialect)]


def _generated_bytes(filename: str, content: str, content_type: str) -> tuple[bytes, str]:
    ext = _extension(filename)
    if ext == ".docx":
        from docx import Document

        doc = Document()
        for block in (content or "").split("\n\n"):
            lines = block.splitlines()
            if not lines:
                doc.add_paragraph("")
            elif len(lines) == 1:
                doc.add_paragraph(lines[0])
            else:
                for line in lines:
                    doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if ext == ".xlsx":
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "Viktor Export"
        for row in _csv_rows(content):
            ws.append(row)
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    return content.encode("utf-8"), content_type or mimetypes.guess_type(filename)[0] or "text/plain; charset=utf-8"


def create_downloadable_file(
    *,
    project_id: str,
    topic_thread_id: str,
    filename: str,
    content: str,
    content_type: str = "text/markdown; charset=utf-8",
) -> dict[str, Any]:
    safe_name = _safe_filename(filename)
    if "." not in safe_name:
        safe_name = f"{safe_name}.md"
    data, detected_type = _generated_bytes(safe_name, content or "", content_type)
    _validate_file(safe_name, len(data))
    object_key = _object_key(project_id, topic_thread_id, safe_name, generated=True)
    _bucket().put_object(object_key, data, headers={"Content-Type": detected_type})
    return {
        "filename": safe_name,
        "size": len(data),
        "content_type": detected_type,
        "oss_uri": _oss_uri(object_key),
        "object_key": object_key,
        "download_url": signed_download_url(object_key),
    }


def format_attachments_for_prompt(attachments: list[dict[str, Any]] | None) -> str:
    if not attachments:
        return ""
    lines = ["\n\n[用户本轮上传的附件]"]
    for idx, item in enumerate(attachments, start=1):
        filename = str(item.get("filename") or "未命名文件")
        size = item.get("size") or 0
        status = str(item.get("extract_status") or "unknown")
        download_url = str(item.get("download_url") or "")
        oss_uri = str(item.get("oss_uri") or "")
        lines.append(f"\n附件 {idx}: {filename}")
        lines.append(f"- size: {size}")
        lines.append(f"- extract_status: {status}")
        if oss_uri:
            lines.append(f"- oss_uri: {oss_uri}")
        if download_url:
            lines.append(f"- download_url: {download_url}")
        text = str(item.get("extracted_text") or item.get("extracted_preview") or "").strip()
        if text:
            lines.append("- readable_content:")
            lines.append("```")
            lines.append(text)
            lines.append("```")
    lines.append("\n请在回答时优先使用附件中的可读内容；如果需要完整原文件，引用 download_url。")
    return "\n".join(lines)


def tool_result_markdown(result: dict[str, Any]) -> str:
    return json.dumps(result, ensure_ascii=False, indent=2)
